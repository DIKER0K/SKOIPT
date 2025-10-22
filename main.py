import os
import re
import json
from collections import defaultdict
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Глобальные переменные для хранения информации о сменах
group_shifts = {}
SHIFTS_FILE = "group_shifts.json"  # Файл для сохранения настроек смен

def parse_schedule_from_docx(file_path):
    """
    Универсальный парсер расписания из DOCX для формата Салаватского колледжа.
    Работает даже с объединёнными ячейками и различным форматированием.
    """
    doc = Document(file_path)
    schedules = {}
    days_ru = ["Понедельник", "Вторник", "Среда", "Четверг", "Пятница", "Суббота"]

    current_group = None
    current_table = None

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        # Находим группу
        group_match = re.search(r'Расписание уроков\s+для\s+(.+?)\s+группы', text)
        if group_match:
            current_group = group_match.group(1).strip()
            schedules[current_group] = {
                'days': {day: {} for day in days_ru},
                'zero_lesson': {day: {} for day in days_ru}
            }
            continue

    # Второй проход — извлекаем все таблицы для найденных групп
    tables = [t for t in doc.tables if any(day in "\n".join(cell.text for row in t.rows for cell in row.cells) for day in days_ru)]

    group_index = 0
    group_names = list(schedules.keys())
    for table in tables:
        if group_index >= len(group_names):
            break
        group_name = group_names[group_index]
        parse_schedule_table_fixed(table, group_name, schedules, days_ru)
        group_index += 1

    return schedules

def parse_schedule_table_fixed(table, group_name, schedules, days_ru):
    """
    Новый, устойчивый парсер таблицы с расписанием.
    Обрабатывает объединённые ячейки и сложное форматирование.
    """
    rows = table.rows
    if not rows:
        return

    # Первая строка — дни недели
    header = [cell.text.strip() for cell in rows[0].cells]
    if len(header) < 2:
        return

    day_columns = {}
    for idx, cell in enumerate(header[1:], 1):
        for day in days_ru:
            if day in cell:
                day_columns[idx] = day
                break

    # Остальные строки — пары
    for r in rows[1:]:
        cells = [c.text.strip() for c in r.cells]
        if not cells or not cells[0]:
            continue

        lesson_num = cells[0].strip()
        if not re.match(r'^\d+$', lesson_num):
            continue  # пропускаем строки без номера пары

        for idx, text in enumerate(cells[1:], 1):
            if idx not in day_columns:
                continue
            day = day_columns[idx]

            if not text.strip():
                continue

            lesson_info = parse_lesson_info_fixed(text)
            if not lesson_info or not lesson_info.get("teacher"):
                continue

            if lesson_num == "0":
                schedules[group_name]['zero_lesson'][day] = lesson_info
            else:
                schedules[group_name]['days'][day][lesson_num] = lesson_info

def parse_lesson_info_fixed(cell_text):
    """
    Усовершенствованный анализ ячейки расписания.
    Учитывает возможные переносы строк, ФИО и аудитории.
    """
    text = re.sub(r'\s+', ' ', cell_text.strip())
    if not text:
        return None

    # Пример строки: "Информатика Ситников Д.А. 221"
    # или "МДК.01.01 Подготовка педагога ... Юльякшина Р.А."
    # или "Физика Сагитова С.Ф."
    teacher_match = re.search(r'([А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.[А-ЯЁ]\.?)', text)
    classroom_match = re.search(r'(\d{2,4}[А-Яа-я]?)\b', text)

    teacher = teacher_match.group(1) if teacher_match else ''
    classroom = classroom_match.group(1) if classroom_match else ''

    # Предмет — всё, что до фамилии преподавателя
    subject = ''
    if teacher_match:
        subject = text[:teacher_match.start()].strip(' ,.;-')
    else:
        # Если ФИО нет — берем всё как предмет
        subject = text.strip()

    # Если предмет остался пустым, но есть кабинет — убираем его из текста
    if not subject and classroom:
        subject = re.sub(r'\b' + re.escape(classroom) + r'\b', '', text).strip()

    # Нормализуем
    subject = re.sub(r'\s+', ' ', subject).strip()
    teacher = teacher.strip()
    classroom = classroom.strip()

    if not subject and not teacher:
        return None

    return {'subject': subject, 'teacher': teacher, 'classroom': classroom}

def parse_schedule_table(table, group_name, schedules, days_ru):
    """
    Парсит таблицу с расписанием
    """
    # Парсим строки таблицы
    for row_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        
        if len(cells) < 2:  # Нужно как минимум номер пары + один день
            continue
            
        # Определяем номер пары (первая ячейка)
        lesson_num = cells[0]
        
        # Пропускаем заголовки и пустые строки
        if not lesson_num or lesson_num in ['', '##', 'Понедельник', 'Вторник', 'Среда', 'Четверг', 'Пятница', 'Суббота']:
            continue
        
        # Обрабатываем каждый день недели
        for day_idx, day in enumerate(days_ru):
            cell_idx = day_idx + 1  # +1 потому что первая колонка - номер пары
            
            if cell_idx < len(cells) and cells[cell_idx].strip() and cells[cell_idx] != '##':
                lesson_info = parse_lesson_info(cells[cell_idx])
                if lesson_info and lesson_info.get('teacher'):
                    if lesson_num == "0":
                        schedules[group_name]['zero_lesson'][day] = lesson_info
                    else:
                        schedules[group_name]['days'][day][lesson_num] = lesson_info

def parse_lesson_info(cell_text):
    """
    Парсит информацию о паре из ячейки таблицы
    """
    # Разбиваем текст на строки и фильтруем пустые
    lines = [line.strip() for line in cell_text.split('\n') if line.strip()]
    
    if not lines:
        return None
    
    lesson_info = {
        'subject': '',
        'teacher': '',
        'classroom': ''
    }
    
    # Первая непустая строка - обычно предмет
    lesson_info['subject'] = lines[0]
    
    # Ищем преподавателя (обычно последняя строка, содержащая ФИО)
    teacher_found = False
    
    # Сначала ищем по строгим паттернам ФИО
    for line in lines:
        clean_line = line.strip()
        # Паттерны для ФИО преподавателей
        patterns = [
            r'^[А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.[А-ЯЁ]\.$',  # Фамилия И.О.
            r'^[А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.[А-ЯЁ]\.?\s*\d*$',  # Фамилия И.О. 214
            r'^[А-ЯЁ][а-яё]+$',  # Просто фамилия
        ]
        
        for pattern in patterns:
            if re.match(pattern, clean_line):
                # Проверяем, что это не название предмета
                subject_indicators = ['МДК', 'УП', 'ПП', 'Основы', 'Математика', 'Физика', 'Химия', 
                                    'История', 'Русский', 'Литература', 'Иностранный', 'Язык',
                                    'Информатика', 'Биология', 'География', 'Физическая', 'Культура']
                if not any(indicator in clean_line for indicator in subject_indicators):
                    lesson_info['teacher'] = clean_line
                    teacher_found = True
                    break
        if teacher_found:
            break
    
    # Если не нашли по паттерну, берем последнюю строку как преподавателя
    if not teacher_found and len(lines) > 1:
        last_line = lines[-1].strip()
        # Проверяем, что последняя строка похожа на ФИО (начинается с заглавной русской буквы)
        if re.match(r'^[А-ЯЁ]', last_line) and len(last_line.split()) <= 3:
            lesson_info['teacher'] = last_line
    
    # Ищем номер кабинета (цифры в конце строк)
    for line in lines:
        classroom_match = re.search(r'(\d{2,4}[А-Яа-я]?)$', line)
        if classroom_match:
            lesson_info['classroom'] = classroom_match.group(1)
            break
    
    return lesson_info

def load_group_shifts():
    global group_shifts
    if os.path.exists(SHIFTS_FILE):
        try:
            with open(SHIFTS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
                # Преобразуем старый формат {"П-311": 1} → {"П-311": {"shift": 1, "room": ""}}
                for k, v in data.items():
                    if isinstance(v, dict):
                        group_shifts[k] = v
                    else:
                        group_shifts[k] = {"shift": v, "room": ""}
            print(f"✅ Загружены настройки для {len(group_shifts)} групп")
            return True
        except Exception as e:
            print(f"❌ Ошибка при загрузке настроек смен: {e}")
    return False


def save_group_shifts():
    """
    Сохраняет настройки смен в файл
    """
    try:
        with open(SHIFTS_FILE, 'w', encoding='utf-8') as f:
            json.dump(group_shifts, f, ensure_ascii=False, indent=2)
        print(f"✅ Настройки смен сохранены в файл {SHIFTS_FILE}")
        return True
    except Exception as e:
        print(f"❌ Ошибка при сохранении настроек смен: {e}")
        return False

def get_group_settings_interactive(groups):
    """
    Интерактивная настройка смен и кабинетов для групп.
    """
    print("\n🎯 НАСТРОЙКА СМЕН И КАБИНЕТОВ ДЛЯ ГРУПП")
    print("=" * 60)

    global group_shifts

    # Показываем текущие настройки
    existing = [g for g in groups if g in group_shifts]
    if existing:
        print("\n📋 Текущие настройки:")
        for g in existing:
            shift_info = group_shifts[g]
            shift = shift_info.get("shift", 1)
            room = shift_info.get("room", "")
            shift_name = {0: "❌ Исключена", 1: "Первая смена", 2: "Вторая смена"}.get(shift, "Неизвестно")
            print(f"  {g}: {shift_name}, кабинет: {room or '—'}")

        change = input("\nИзменить существующие настройки? (y/n): ").strip().lower()
        if change != 'y':
            print("✅ Используются сохранённые настройки.")
            return group_shifts

    for group in groups:
        info = group_shifts.get(group, {"shift": 1, "room": ""})
        current_shift = info.get("shift", 1)
        current_room = info.get("room", "")

        print(f"\nГруппа: {group}")
        print("1 - Первая смена")
        print("2 - Вторая смена")
        print("3 - Не включать в расписание преподавателей")
        print("0 - Пропустить (оставить текущую)")

        while True:
            choice = input(f"Выберите смену (0/1/2/3) [текущая: {current_shift}]: ").strip()
            if choice == '0':
                break
            elif choice == '1':
                current_shift = 1
                break
            elif choice == '2':
                current_shift = 2
                break
            elif choice == '3':
                current_shift = 0
                break
            else:
                print("❌ Неверный выбор. Попробуйте снова.")

        # Ввод кабинета
        if current_shift != 0:  # если группа не исключена
            room = input(f"Введите кабинет для группы {group} [текущий: {current_room or '—'}]: ").strip()
            if room:
                current_room = room

        group_shifts[group] = {"shift": current_shift, "room": current_room}

    save_group_shifts()
    print("\n✅ Настройки сохранены!")
    return group_shifts



def normalize_teacher_name(name):
    """
    Очищает имя преподавателя от предметов и лишнего текста.
    Делает формат единообразным: Фамилия И.О.
    """
    if not name:
        return None

    # Убираем лишние пробелы
    name = re.sub(r'\s+', ' ', name.strip())

    # Убираем предметы и служебные слова (если "прилипли" к ФИО)
    # Например: "МДК.01.01 Информатика Хусаинова Л.Н." -> "Хусаинова Л.Н."
    match = re.search(r'([А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.[А-ЯЁ]\.?)', name)
    if match:
        name = match.group(1)

    # Убираем кабинеты
    name = re.sub(r'\s*\d{2,4}[А-Яа-я]?$', '', name)

    # Стандартизируем инициалы
    name = re.sub(r'([А-ЯЁ])\.([А-ЯЁ])$', r'\1.\2.', name)
    name = re.sub(r'\.\.', '.', name)
    name = re.sub(r'\s+', ' ', name).strip()

    # Иногда остаются хвосты вроде "Арзамасова А.В.)" или "Арзамасова А.В.-"
    name = re.sub(r'[^А-Яа-яЁё.\s-]', '', name)
    name = re.sub(r'\s+$', '', name)

    # Проверяем, что это похоже на ФИО
    if not re.match(r'^[А-ЯЁ][а-яё]+\s+[А-ЯЁ]\.[А-ЯЁ]\.?$', name):
        return None

    return name

def create_teacher_schedules(student_schedules):
    teacher_schedules = defaultdict(lambda: {
        'first_shift': defaultdict(dict),
        'second_shift': defaultdict(dict)
    })
    
    for group, schedule_data in student_schedules.items():
        info = group_shifts.get(group, {"shift": 1, "room": ""})
        shift = info.get("shift", 1)
        room_default = info.get("room", "")
        if shift == 0:
            continue  # ❌ пропускаем исключённые группы
        
        shift_key = 'first_shift' if shift == 1 else 'second_shift'
        
        # Нулевые пары
        for day, zero_lesson in schedule_data['zero_lesson'].items():
            if zero_lesson and zero_lesson.get('teacher'):
                teacher = normalize_teacher_name(zero_lesson['teacher'])
                if teacher:
                    classroom = zero_lesson['classroom'] or room_default
                    teacher_schedules[teacher][shift_key][day]['0'] = {
                        'subject': zero_lesson['subject'],
                        'group': group,
                        'classroom': classroom
                    }
        
        # Основные пары
        for day, lessons in schedule_data['days'].items():
            for lesson_num, lesson_info in lessons.items():
                if lesson_info and lesson_info.get('teacher'):
                    teacher = normalize_teacher_name(lesson_info['teacher'])
                    if teacher:
                        classroom = lesson_info['classroom'] or room_default
                        teacher_schedules[teacher][shift_key][day][lesson_num] = {
                            'subject': lesson_info['subject'],
                            'group': group,
                            'classroom': classroom
                        }
    
    return dict(teacher_schedules)



def create_teacher_schedule_docx(teacher_name, schedule, output_folder):
    """
    Создает DOCX файл с расписанием преподавателя.
    Вторая смена — с новой страницы.
    """
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Заголовок
    title = doc.add_heading('Расписание учебных занятий', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    teacher_para = doc.add_paragraph()
    run = teacher_para.add_run(f"Преподаватель: {teacher_name}")
    run.bold = True
    run.font.size = Pt(14)
    teacher_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    period_para = doc.add_paragraph("Период: с 20.10.2025 г.")
    period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    days_ru = ["Понедельник", "Вторник", "Среда", "Четверг", "Пятница", "Суббота"]
    lesson_numbers = ['0', '1', '2', '3', '4', '5']

    def fill_table(shift_data, title_text):
        if not any(shift_data.values()):
            return
        doc.add_paragraph().add_run(title_text).bold = True
        doc.add_paragraph()
        table = doc.add_table(rows=len(lesson_numbers) + 1, cols=len(days_ru) + 1)
        table.style = 'Table Grid'

        hdr = table.rows[0].cells
        hdr[0].text = "№ пары"
        hdr[0].paragraphs[0].runs[0].bold = True
        for i, d in enumerate(days_ru):
            hdr[i + 1].text = d
            hdr[i + 1].paragraphs[0].runs[0].bold = True

        for r, num in enumerate(lesson_numbers):
            row = table.rows[r + 1].cells
            row[0].text = num
            row[0].paragraphs[0].runs[0].bold = True
            for c, day in enumerate(days_ru):
                if day in shift_data and num in shift_data[day]:
                    info = shift_data[day][num]
                    subject = info.get('subject', '')
                    group = info.get('group', '')
                    classroom = info.get('classroom', '')
                    text = ''
                    if subject:
                        text += f"{subject}\n"
                    if group:
                        text += f"Группа: {group}\n"
                    if classroom:
                        text += f"Ауд.: {classroom}"
                    row[c + 1].text = text.strip()

    # первая смена
    fill_table(schedule['first_shift'], "ПЕРВАЯ СМЕНА")

    # вторая смена — с новой страницы
    if any(schedule['second_shift'].values()):
        doc.add_page_break()
        fill_table(schedule['second_shift'], "ВТОРАЯ СМЕНА")

    if not any(schedule['first_shift'].values()) and not any(schedule['second_shift'].values()):
        doc.add_paragraph("Нет учебных занятий в расписании").alignment = WD_ALIGN_PARAGRAPH.CENTER

    safe_name = re.sub(r'[<>:"/\\|?*]', '_', teacher_name)
    filepath = os.path.join(output_folder, f"Расписание_{safe_name}.docx")
    doc.save(filepath)
    return filepath

def print_teacher_statistics(teacher_schedules):
    """
    Показывает статистику по каждому преподавателю: количество групп, предметов и пар.
    """
    print("\n📊 СТАТИСТИКА ПО ПРЕПОДАВАТЕЛЯМ")
    print("=" * 60)

    total_lessons = 0
    total_teachers = len(teacher_schedules)

    for teacher, data in sorted(teacher_schedules.items()):
        groups = set()
        subjects = set()
        lessons_count = 0

        for shift_key in ['first_shift', 'second_shift']:
            for day, lessons in data[shift_key].items():
                for num, info in lessons.items():
                    groups.add(info['group'])
                    subjects.add(info['subject'])
                    lessons_count += 1

        total_lessons += lessons_count

        print(f"👨‍🏫 {teacher}: {lessons_count} пар, {len(groups)} групп, {len(subjects)} предметов")

    print("=" * 60)
    print(f"Всего преподавателей: {total_teachers}")
    print(f"Всего пар: {total_lessons}")


def debug_schedule_content(schedules):
    """Выводит отладочную информацию о расписаниях"""
    print("\n=== ОТЛАДОЧНАЯ ИНФОРМАЦИЯ О РАСПИСАНИЯХ ===")
    
    total_lessons = 0
    total_teachers = set()
    
    for group_name, schedule_data in schedules.items():
        group_lessons = 0
        group_teachers = set()
        
        # Считаем обычные пары
        for day, lessons in schedule_data['days'].items():
            for lesson_num, lesson_info in lessons.items():
                if lesson_info and lesson_info.get('teacher'):
                    group_lessons += 1
                    group_teachers.add(lesson_info['teacher'])
                    total_teachers.add(lesson_info['teacher'])
        
        # Считаем нулевые пары
        for day, zero_lesson in schedule_data['zero_lesson'].items():
            if zero_lesson and zero_lesson.get('teacher'):
                group_lessons += 1
                group_teachers.add(zero_lesson['teacher'])
                total_teachers.add(zero_lesson['teacher'])
        
        total_lessons += group_lessons
        
        if group_lessons > 0:
            print(f"  {group_name}: {group_lessons} пар, {len(group_teachers)} преподавателей")
        else:
            print(f"  {group_name}: НЕТ ДАННЫХ")
    
    print(f"\n📊 ИТОГО: {total_lessons} пар, {len(total_teachers)} уникальных преподавателей")

def main():
    """
    Основная функция
    """
    input_file = "Расписание.docx"
    output_dir = "Расписания_преподавателей"
    
    # Проверяем существование файла
    if not os.path.exists(input_file):
        print(f"❌ Файл '{input_file}' не найден!")
        print("📁 Файлы в текущей директории:")
        for file in os.listdir('.'):
            if file.endswith('.docx'):
                print(f"  - {file}")
        return
    
    # Создаем папку для выходных файлов
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    print("🔍 Парсим расписание для студентов...")
    
    try:
        student_schedules = parse_schedule_from_docx(input_file)
        
        if not student_schedules:
            print("❌ Не удалось найти расписания в файле")
            return
        
        print(f"✅ Найдено групп: {len(student_schedules)}")
        
        # Отладочная информация
        debug_schedule_content(student_schedules)
        
        # Загружаем сохраненные настройки смен
        loaded = load_group_shifts()
        
        # Получаем информацию о сменах
        groups = list(student_schedules.keys())
        get_group_settings_interactive(groups)
        
        print("👨‍🏫 Создаем расписания для преподавателей...")
        teacher_schedules = create_teacher_schedules(student_schedules)
        print_teacher_statistics(teacher_schedules)
        
        print(f"✅ Найдено преподавателей: {len(teacher_schedules)}")
        print("📋 Список преподавателей:")
        for i, teacher in enumerate(sorted(teacher_schedules.keys()), 1):
            print(f"  {i:2d}. {teacher}")
        
        # Создаем файлы для каждого преподавателя
        created_files = []
        for teacher_name, schedule in teacher_schedules.items():
            if teacher_name and teacher_name.strip() and teacher_name != "Обществознание":
                filepath = create_teacher_schedule_docx(teacher_name, schedule, output_dir)
                created_files.append(filepath)
                print(f"📄 Создан файл: {os.path.basename(filepath)}")
        
        print(f"\n🎉 Готово! Создано {len(created_files)} файлов в папке '{output_dir}'")
        
        # Показываем статистику по сменам
        print("\n📊 СТАТИСТИКА ПО СМЕНАМ:")
        first_shift_groups = [group for group, shift in group_shifts.items() if shift == 1]
        second_shift_groups = [group for group, shift in group_shifts.items() if shift == 2]
        print(f"   Первая смена: {len(first_shift_groups)} групп")
        print(f"   Вторая смена: {len(second_shift_groups)} групп")
        
        # Сохраняем финальные настройки
        save_group_shifts()
        
    except Exception as e:
        print(f"❌ Произошла ошибка: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()